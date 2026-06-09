#!/usr/bin/env python3
"""Export raw policy thoughts and command JSON from prompted-run traces.

The runtime writes ``policy_raw_*.jsonl`` rows when ``QARM_POLICY_TRACE_SAVE_RAW``
is enabled. This tool turns those rows into one compact JSON document that is
easy to line up with manually captured run photos.
"""

from __future__ import annotations

import argparse
import csv
import json
import re
import sys
from datetime import datetime, timezone
from pathlib import Path
from typing import Any


TRACE_GLOB = "policy_raw_*.jsonl"
DEFAULT_PROMPTS = "1,2,3"


def project_root() -> Path:
    return Path(__file__).resolve().parents[1]


def utc_now_iso() -> str:
    return datetime.now(timezone.utc).isoformat(timespec="seconds")


def timestamp_ms_to_iso(timestamp_ms: object, *, local: bool) -> str:
    try:
        seconds = int(timestamp_ms) / 1000.0
    except (TypeError, ValueError):
        return ""
    dt = datetime.fromtimestamp(seconds, timezone.utc)
    if local:
        dt = dt.astimezone()
    return dt.isoformat(timespec="seconds")


def read_json_file(path: Path) -> Any:
    try:
        with path.open("r", encoding="utf-8") as fp:
            return json.load(fp)
    except FileNotFoundError as exc:
        raise SystemExit(f"ERROR: file not found: {path}") from exc
    except json.JSONDecodeError as exc:
        raise SystemExit(f"ERROR: invalid JSON in {path}: {exc}") from exc


def load_prompt_index(prompts_path: Path) -> dict[str, dict[str, Any]]:
    if not prompts_path.exists():
        return {}
    payload = read_json_file(prompts_path)
    prompts = payload.get("prompts") if isinstance(payload, dict) else payload
    if not isinstance(prompts, list):
        return {}
    out: dict[str, dict[str, Any]] = {}
    for idx, prompt in enumerate(prompts, start=1):
        if not isinstance(prompt, dict):
            continue
        prompt_id = str(prompt.get("id", "")).strip()
        if not prompt_id:
            continue
        row = dict(prompt)
        row.setdefault("prompt_index", idx)
        row.setdefault("prompt_text", row.get("text", ""))
        out[prompt_id] = row
    return out


def parse_prompt_filter(raw: str) -> set[str]:
    values: set[str] = set()
    for part in str(raw or "").replace(";", ",").split(","):
        item = part.strip().lower()
        if item:
            values.add(item)
    return values


def infer_prompt_index(prompt_id: str) -> str:
    match = re.search(r"(?:^|_)prompt[_-]?(\d+)(?:_|$)", str(prompt_id).lower())
    if match:
        return match.group(1)
    match = re.search(r"^p(\d+)$", str(prompt_id).lower())
    if match:
        return match.group(1)
    return ""


def normalize_prompt_index(value: object, prompt_id: str, prompt_index_by_id: dict[str, dict[str, Any]]) -> str:
    raw = str(value if value is not None else "").strip()
    if raw:
        return raw
    prompt_row = prompt_index_by_id.get(str(prompt_id))
    if prompt_row is not None:
        candidate = str(prompt_row.get("prompt_index", "")).strip()
        if candidate:
            return candidate
    return infer_prompt_index(prompt_id)


def prompt_matches(
    metadata: dict[str, Any],
    prompt_filters: set[str],
    prompt_index_by_id: dict[str, dict[str, Any]],
) -> bool:
    if not prompt_filters:
        return True
    prompt_id = str(metadata.get("prompt_id", "")).strip()
    prompt_index = normalize_prompt_index(metadata.get("prompt_index", ""), prompt_id, prompt_index_by_id)
    aliases = {
        prompt_id.lower(),
        prompt_index.lower(),
        f"p{prompt_index}".lower(),
        f"prompt_{prompt_index}".lower(),
    }
    if prompt_id.lower().startswith("prompt_"):
        parts = prompt_id.lower().split("_")
        if len(parts) > 1 and parts[1].isdigit():
            aliases.add(parts[1])
            aliases.add(f"p{parts[1]}")
            aliases.add(f"prompt_{parts[1]}")
    return bool(prompt_filters & aliases)


def read_csv_rows(path: Path) -> list[dict[str, str]]:
    if not path.exists():
        return []
    with path.open("r", encoding="utf-8", newline="") as fp:
        return [dict(row) for row in csv.DictReader(fp)]


def csv_metadata_for_trial(trial_dir: Path) -> dict[str, str]:
    session_csv = trial_dir.parent / "trials.csv"
    trial_id = trial_dir.name
    for row in read_csv_rows(session_csv):
        if str(row.get("trial_id", "")).strip() == trial_id:
            return row
    return {}


def read_trial_meta(trial_dir: Path) -> dict[str, Any]:
    meta_path = trial_dir / "trial_meta.json"
    if not meta_path.exists():
        return {}
    payload = read_json_file(meta_path)
    if not isinstance(payload, dict):
        raise SystemExit(f"ERROR: {meta_path} must contain a JSON object")
    return payload


def discover_trace_files(paths: list[Path]) -> list[Path]:
    traces: set[Path] = set()
    for raw_path in paths:
        path = raw_path.resolve()
        if path.is_file():
            if path.name.startswith("policy_raw_") and path.suffix == ".jsonl":
                traces.add(path)
            continue
        if not path.exists():
            raise SystemExit(f"ERROR: path does not exist: {path}")
        for trace_path in path.rglob(TRACE_GLOB):
            if trace_path.is_file():
                traces.add(trace_path.resolve())
    return sorted(traces)


def group_traces_by_trial(trace_files: list[Path]) -> dict[Path, list[Path]]:
    grouped: dict[Path, list[Path]] = {}
    for trace_file in trace_files:
        grouped.setdefault(trace_file.parent.resolve(), []).append(trace_file)
    return {trial_dir: sorted(paths) for trial_dir, paths in sorted(grouped.items(), key=lambda item: str(item[0]))}


def trial_metadata(
    trial_dir: Path,
    *,
    prompt_index_by_id: dict[str, dict[str, Any]],
) -> dict[str, Any]:
    meta = read_trial_meta(trial_dir)
    csv_row = csv_metadata_for_trial(trial_dir)
    merged: dict[str, Any] = {}
    merged.update(meta)
    merged.update({key: value for key, value in csv_row.items() if value not in (None, "")})

    prompt_id = str(merged.get("prompt_id", "")).strip()
    prompt_row = prompt_index_by_id.get(prompt_id, {})
    if not merged.get("prompt_index"):
        merged["prompt_index"] = normalize_prompt_index("", prompt_id, prompt_index_by_id)
    if not merged.get("prompt_text"):
        merged["prompt_text"] = prompt_row.get("text", prompt_row.get("prompt_text", ""))
    if not merged.get("trial_id"):
        merged["trial_id"] = trial_dir.name
    merged["trial_dir"] = str(trial_dir)
    merged["session"] = trial_dir.parent.name
    return merged


def success_source(metadata: dict[str, Any], *, strict_success_csv: bool) -> str:
    final_result = str(metadata.get("final_result", "")).strip().lower()
    if final_result == "success":
        return "csv_final_result"
    if strict_success_csv:
        return ""
    status = str(metadata.get("status", "")).strip().lower()
    return_code = str(metadata.get("return_code", "")).strip()
    crashed = str(metadata.get("crashed", "")).strip().lower()
    if status == "finished" and return_code in {"", "0"} and crashed not in {"true", "1", "yes"}:
        return "meta_finished_return_code"
    return ""


def parse_first_json_object(text: str) -> tuple[dict[str, Any] | None, int, int]:
    decoder = json.JSONDecoder()
    for match in re.finditer(r"\{", text):
        start = match.start()
        try:
            payload, consumed = decoder.raw_decode(text[start:])
        except json.JSONDecodeError:
            continue
        if isinstance(payload, dict):
            return payload, start, start + consumed
    return None, -1, -1


def fallback_decision_from_text(text: str) -> dict[str, str]:
    command_match = re.search(r'"command"\s*:\s*"(?P<command>[^"]+)"', text)
    reason_match = re.search(r'"reason"\s*:\s*"(?P<reason>[^"]*)"', text)
    out: dict[str, str] = {}
    if command_match:
        out["command"] = command_match.group("command").strip()
    if reason_match:
        out["reason"] = reason_match.group("reason").strip()
    return out


def parse_raw_policy_output(raw_output: object) -> dict[str, Any]:
    text = str(raw_output or "")
    start_tag = "FINAL_JSON_START"
    end_tag = "FINAL_JSON_END"
    start_index = text.find(start_tag)
    end_index = text.find(end_tag, start_index + len(start_tag)) if start_index >= 0 else -1

    thought = ""
    policy_output_text = ""
    payload: dict[str, Any] | None = None
    parse_error = ""

    if start_index >= 0 and end_index >= 0:
        thought = text[:start_index].strip()
        block = text[start_index + len(start_tag) : end_index].strip()
        payload, payload_start, payload_end = parse_first_json_object(block)
        if payload is not None:
            policy_output_text = block[payload_start:payload_end].strip()
        else:
            policy_output_text = block
            parse_error = "final_json_block_contains_no_json_object"
    else:
        payload, payload_start, payload_end = parse_first_json_object(text)
        if payload is not None:
            thought = text[:payload_start].strip()
            policy_output_text = text[payload_start:payload_end].strip()
        else:
            thought = text.strip()
            parse_error = "raw_output_contains_no_json_object"

    if payload is None:
        payload = fallback_decision_from_text(text)
    command = str(payload.get("command", "")).strip() if isinstance(payload, dict) else ""
    reason = str(payload.get("reason", "")).strip() if isinstance(payload, dict) else ""

    return {
        "command": command,
        "reason": reason,
        "raw_policy_thought": thought,
        "policy_output": payload if isinstance(payload, dict) else {},
        "policy_output_text": policy_output_text,
        "parse_error": parse_error,
    }


def read_policy_rows(trace_path: Path) -> list[tuple[int, dict[str, Any]]]:
    rows: list[tuple[int, dict[str, Any]]] = []
    with trace_path.open("r", encoding="utf-8") as fp:
        for line_number, line in enumerate(fp, start=1):
            text = line.strip()
            if not text:
                continue
            try:
                row = json.loads(text)
            except json.JSONDecodeError as exc:
                raise SystemExit(f"ERROR: invalid JSONL in {trace_path}:{line_number}: {exc}") from exc
            if not isinstance(row, dict):
                raise SystemExit(f"ERROR: {trace_path}:{line_number} is not a JSON object")
            rows.append((line_number, row))
    return rows


def extract_commands(trace_files: list[Path]) -> list[dict[str, Any]]:
    commands: list[dict[str, Any]] = []
    command_index = 0
    for trace_file in trace_files:
        for line_number, row in read_policy_rows(trace_file):
            parsed = parse_raw_policy_output(row.get("raw_output", ""))
            command = parsed.get("command", "")
            if not command:
                raise SystemExit(
                    f"ERROR: could not extract command from {trace_file}:{line_number}; "
                    f"parse_error={parsed.get('parse_error') or 'missing_command'}"
                )
            command_index += 1
            llm_input = row.get("llm_input") if isinstance(row.get("llm_input"), dict) else {}
            state = llm_input.get("state", {}) if isinstance(llm_input, dict) else {}
            allowed_commands = llm_input.get("allowed_commands", []) if isinstance(llm_input, dict) else []
            commands.append(
                {
                    "photo_match_index": command_index,
                    "command_index": command_index,
                    "timestamp_ms": row.get("timestamp_ms", ""),
                    "timestamp_utc": timestamp_ms_to_iso(row.get("timestamp_ms", ""), local=False),
                    "timestamp_local": timestamp_ms_to_iso(row.get("timestamp_ms", ""), local=True),
                    "cycle": row.get("cycle", ""),
                    "step_index": row.get("step_index", ""),
                    "phase": row.get("phase", ""),
                    "allowed_commands": allowed_commands,
                    "policy_input_state": state,
                    "command": command,
                    "reason": parsed.get("reason", ""),
                    "policy_output": parsed.get("policy_output", {}),
                    "policy_output_text": parsed.get("policy_output_text", ""),
                    "raw_policy_thought": parsed.get("raw_policy_thought", ""),
                    "raw_policy_output": row.get("raw_output", ""),
                    "parse_error": parsed.get("parse_error", ""),
                    "trace_file": str(trace_file),
                    "trace_line": line_number,
                }
            )
    return commands


def build_run_export(
    trial_dir: Path,
    trace_files: list[Path],
    metadata: dict[str, Any],
    success_kind: str,
) -> dict[str, Any]:
    commands = extract_commands(trace_files)
    return {
        "session": metadata.get("session", ""),
        "trial_id": metadata.get("trial_id", trial_dir.name),
        "trial_dir": str(trial_dir),
        "prompt_id": metadata.get("prompt_id", ""),
        "prompt_index": metadata.get("prompt_index", ""),
        "repeat_index": metadata.get("repeat_index", ""),
        "final_result": metadata.get("final_result", ""),
        "success_source": success_kind,
        "autonomous": metadata.get("autonomous", ""),
        "outcome_alias": metadata.get("outcome_alias", ""),
        "started_at": metadata.get("started_at", ""),
        "ended_at": metadata.get("ended_at", ""),
        "duration_s": metadata.get("duration_s", ""),
        "initial_condition": metadata.get("initial_condition", ""),
        "prompt_text": metadata.get("prompt_text", ""),
        "console_log": metadata.get("console_log", str(trial_dir / "console_log.txt")),
        "policy_trace_files": [str(path) for path in trace_files],
        "command_count": len(commands),
        "commands": commands,
    }


def parse_docx_run_labels(raw: str) -> dict[str, str]:
    labels: dict[str, str] = {}
    for part in re.split(r"[;,]", str(raw or "")):
        item = part.strip()
        if not item:
            continue
        if "=" not in item:
            raise SystemExit(f"ERROR: invalid --docx-run-labels item {item!r}; use key=value")
        key, value = item.split("=", 1)
        key = key.strip()
        value = value.strip()
        if not key or not value:
            raise SystemExit(f"ERROR: invalid --docx-run-labels item {item!r}; use key=value")
        labels[key] = value
    return labels


def docx_run_label(run: dict[str, Any], labels: dict[str, str]) -> str:
    for key in (
        str(run.get("trial_id", "")).strip(),
        str(run.get("repeat_index", "")).strip(),
        str(run.get("prompt_id", "")).strip(),
    ):
        if key and key in labels:
            return labels[key]
    prompt_index = str(run.get("prompt_index", "")).strip()
    prompt_id = str(run.get("prompt_id", "")).strip()
    trial_id = str(run.get("trial_id", "")).strip()
    repeat_index = str(run.get("repeat_index", "")).strip()
    parts = []
    if prompt_index:
        parts.append(f"Prompt {prompt_index}")
    if prompt_id:
        parts.append(prompt_id)
    if trial_id:
        parts.append(trial_id)
    if repeat_index:
        parts.append(f"repeat {repeat_index}")
    return " - ".join(parts) if parts else "Policy run"


def command_docx_text(command: dict[str, Any], field: str) -> list[tuple[str, str]]:
    reason = str(command.get("reason", "") or command.get("policy_output", {}).get("reason", "")).strip()
    thought = str(command.get("raw_policy_thought", "")).strip()
    if field == "reason":
        return [("Policy reason", reason or "No policy reason captured.")]
    if field == "thought":
        return [("Raw policy thought", thought or "No raw policy thought captured.")]
    return [
        ("Policy reason", reason or "No policy reason captured."),
        ("Raw policy thought", thought or "No raw policy thought captured."),
    ]


def write_docx_export(
    payload: dict[str, Any],
    output_path: Path,
    *,
    title: str,
    field: str,
    run_labels: dict[str, str],
) -> None:
    try:
        from docx import Document
        from docx.enum.section import WD_SECTION
        from docx.enum.text import WD_BREAK, WD_LINE_SPACING
        from docx.shared import Inches, Pt, RGBColor
    except ImportError as exc:
        raise SystemExit(
            "ERROR: --docx-output requires python-docx. "
            "Run with the bundled workspace Python runtime if system Python lacks it."
        ) from exc

    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)

    blue = RGBColor(46, 116, 181)
    dark_blue = RGBColor(31, 77, 120)
    muted = RGBColor(89, 89, 89)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE

    for style_name, size, color, before, after in [
        ("Title", 20, blue, 0, 8),
        ("Subtitle", 11, muted, 0, 12),
        ("Heading 1", 16, blue, 18, 10),
        ("Heading 2", 13, blue, 14, 7),
        ("Heading 3", 12, dark_blue, 10, 5),
    ]:
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.25
        style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE

    doc.add_paragraph(title, style="Title")
    subtitle = doc.add_paragraph(style="Subtitle")
    subtitle.add_run(
        f"Runs: {payload.get('run_count', 0)} | Commands: {payload.get('command_count', 0)} | Field: {field}"
    )

    runs = list(payload.get("runs", []))
    for run_index, run in enumerate(runs):
        if run_index > 0:
            doc.add_section(WD_SECTION.NEW_PAGE)
        doc.add_paragraph(docx_run_label(run, run_labels), style="Heading 1")

        meta = doc.add_paragraph()
        for label, value in [
            ("Trial", run.get("trial_id", "")),
            ("Repeat", run.get("repeat_index", "")),
            ("Result", run.get("final_result", "")),
            ("Outcome", run.get("outcome_alias", "")),
            ("Initial condition", run.get("initial_condition", "")),
        ]:
            if value not in (None, ""):
                label_run = meta.add_run(f"{label}: ")
                label_run.bold = True
                meta.add_run(f"{value}  ")

        mission_text = str(run.get("prompt_text", "") or "").strip()
        if mission_text:
            doc.add_paragraph("Mission Prompt", style="Heading 2")
            mission = doc.add_paragraph()
            for idx, line in enumerate(mission_text.splitlines()):
                if idx:
                    mission.add_run().add_break(WD_BREAK.LINE)
                mission.add_run(line)

        doc.add_paragraph("Command List", style="Heading 2")
        for command in run.get("commands", []):
            command_index = command.get("command_index") or command.get("photo_match_index") or ""
            command_name = str(command.get("command", "")).strip() or "unknown"
            heading = doc.add_paragraph(style="Heading 3")
            heading.add_run(f"{command_index}. Command: ").bold = True
            heading.add_run(command_name).bold = True

            for label, text in command_docx_text(command, field):
                paragraph = doc.add_paragraph()
                paragraph.paragraph_format.left_indent = Inches(0.25)
                label_run = paragraph.add_run(f"{label}: ")
                label_run.bold = True
                label_run.font.color.rgb = muted
                paragraph.add_run(text)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)


def parse_args(argv: list[str]) -> argparse.Namespace:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument(
        "paths",
        nargs="*",
        default=["results"],
        help="Results/session/trial directories or policy_raw_*.jsonl files. Defaults to results/.",
    )
    parser.add_argument(
        "--prompts",
        default=DEFAULT_PROMPTS,
        help="Comma-separated prompt filters by index/id/alias. Default: 1,2,3.",
    )
    parser.add_argument(
        "--prompts-file",
        default="experiments/prompts.json",
        help="Prompt metadata JSON used to resolve prompt indexes. Default: experiments/prompts.json.",
    )
    parser.add_argument(
        "--include-non-success",
        action="store_true",
        help="Include matching runs even when final_result is not success.",
    )
    parser.add_argument(
        "--strict-success-csv",
        action="store_true",
        help="For successful-only mode, require trials.csv final_result=success; do not infer from finished trial_meta.json.",
    )
    parser.add_argument(
        "--output",
        "-o",
        default="",
        help="Write JSON to this file instead of stdout.",
    )
    parser.add_argument(
        "--docx-output",
        default="",
        help="Optional DOCX summary output path.",
    )
    parser.add_argument(
        "--docx-field",
        choices=["reason", "thought", "reason_and_thought"],
        default="reason",
        help="Text field to include in DOCX command entries. Default: reason.",
    )
    parser.add_argument(
        "--docx-title",
        default="Policy Command Summary",
        help="Title for --docx-output.",
    )
    parser.add_argument(
        "--docx-run-labels",
        default="",
        help="Optional labels such as 'trial_0005=Prompt 4A;trial_0010=Prompt 4B'.",
    )
    return parser.parse_args(argv)


def main(argv: list[str] | None = None) -> int:
    args = parse_args(sys.argv[1:] if argv is None else argv)
    root = project_root()
    input_paths = [Path(path) if Path(path).is_absolute() else root / path for path in args.paths]
    prompts_path = Path(args.prompts_file)
    if not prompts_path.is_absolute():
        prompts_path = root / prompts_path
    prompt_index_by_id = load_prompt_index(prompts_path)
    prompt_filters = parse_prompt_filter(args.prompts)

    trace_files = discover_trace_files(input_paths)
    if not trace_files:
        raise SystemExit(f"ERROR: no {TRACE_GLOB} files found under: {', '.join(str(path) for path in input_paths)}")

    runs: list[dict[str, Any]] = []
    skipped: dict[str, int] = {
        "prompt_filter": 0,
        "non_success": 0,
        "empty_trace": 0,
    }
    for trial_dir, grouped_trace_files in group_traces_by_trial(trace_files).items():
        metadata = trial_metadata(trial_dir, prompt_index_by_id=prompt_index_by_id)
        if not prompt_matches(metadata, prompt_filters, prompt_index_by_id):
            skipped["prompt_filter"] += 1
            continue
        success_kind = success_source(metadata, strict_success_csv=bool(args.strict_success_csv))
        if not args.include_non_success and not success_kind:
            skipped["non_success"] += 1
            continue
        if sum(path.stat().st_size for path in grouped_trace_files) <= 0:
            skipped["empty_trace"] += 1
            continue
        runs.append(
            build_run_export(
                trial_dir=trial_dir,
                trace_files=grouped_trace_files,
                metadata=metadata,
                success_kind=success_kind or "included_non_success",
            )
        )

    if not runs:
        raise SystemExit(
            "ERROR: no matching policy trace runs found "
            f"(skipped={skipped}, prompts={sorted(prompt_filters) or 'all'})."
        )

    payload = {
        "schema": "policy_thought_export_v1",
        "generated_at": utc_now_iso(),
        "source_paths": [str(path.resolve()) for path in input_paths],
        "filters": {
            "prompts": sorted(prompt_filters),
            "include_non_success": bool(args.include_non_success),
            "strict_success_csv": bool(args.strict_success_csv),
        },
        "run_count": len(runs),
        "command_count": sum(int(run.get("command_count", 0)) for run in runs),
        "skipped": skipped,
        "runs": runs,
    }

    output_text = json.dumps(payload, indent=2, ensure_ascii=False)
    if args.output:
        output_path = Path(args.output)
        if not output_path.is_absolute():
            output_path = root / output_path
        output_path.parent.mkdir(parents=True, exist_ok=True)
        with output_path.open("w", encoding="utf-8") as fp:
            fp.write(output_text)
            fp.write("\n")
        print(
            f"Wrote {len(runs)} runs / {payload['command_count']} commands to {output_path}",
            file=sys.stderr,
        )
    else:
        print(output_text)
    if args.docx_output:
        docx_output_path = Path(args.docx_output)
        if not docx_output_path.is_absolute():
            docx_output_path = root / docx_output_path
        write_docx_export(
            payload,
            docx_output_path,
            title=str(args.docx_title),
            field=str(args.docx_field),
            run_labels=parse_docx_run_labels(str(args.docx_run_labels)),
        )
        print(f"Wrote DOCX summary to {docx_output_path}", file=sys.stderr)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
